"""Generate body power consumption report — 3 scenarios × sub-states."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

s = doc.sections[0]
for m in ['top', 'bottom', 'left', 'right']:
    setattr(s, f'{m}_margin', Cm(2))

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text); run.bold = True; run.font.size = Pt(9)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shading = c._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'D9D9D9'
        })
        shading.append(shd)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]; cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text)); run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══ TITLE ═══
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('机身功耗测试数据')
r.bold = True; r.font.size = Pt(18)
r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('peripheral_server_sleep 工程实测 | 2026-08-10')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(128, 128, 128)
doc.add_paragraph()

# ═══ 测试条件 ═══
heading('一、测试条件', level=2)
make_table(
    ['参数', '值'],
    [
        ['芯片', 'RSL10'],
        ['供电', 'LDO，VCC=1.10V (0dBm)'],
        ['BLE 低功耗时钟', '8 MHz + BB_DEEP_SLEEP'],
        ['RM 工作时钟', '16 MHz + BB_WAKEUP + CPCLK_PRESCALE_8'],
        ['BLE 广播间隔', '100ms（常规）/ 200ms（轮询）'],
        ['TX Power', '0 dBm'],
    ],
    col_widths=[5, 11]
)

# ═══ 场景1 ═══
heading('二、场景1：纯 BLE 从机', level=2)
para('机身仅作为 BLE Peripheral，不开启 RM。根据连接对象分为三个子状态。')

make_table(
    ['子状态', '功耗', '备注'],
    [
        ['无手机连接，无主机连接（仅广播）', '600 μA', 'deep sleep 正常，100ms 广播间隔'],
        ['手机连接，无主机连接', '600 μA', 'deep sleep 正常，20ms 连接间隔，latency=10'],
        ['无手机连接，主机连接（仅对侧耳）', '600 μA', 'deep sleep 正常，500ms 连接间隔，latency=10'],
        ['手机连接，主机连接（都连上了）', '600 μA', 'deep sleep 正常，手机 20ms + 主机 500ms'],
    ],
    col_widths=[5.5, 2.5, 8]
)

# ═══ 场景2 ═══
heading('三、场景2：BLE 从机 + RM 从机', level=2)
para('机身作为 BLE Peripheral 的同时开启 RM RX。根据 RM 和 BLE 的连接状态分为四个子状态。')

make_table(
    ['子状态', '功耗', '备注'],
    [
        ['RM 未连接，BLE 未连接', '850 μA', 'RM 搜索中 + BLE 仅广播'],
        ['RM 未连接，BLE 连接', '850 μA', 'RM 搜索中 + BLE 保持手机/主机连接'],
        ['RM 连接，BLE 未连接', '1 mA', 'RM 推流接收 + BLE 仅广播'],
        ['RM 连接，BLE 连接', '1 mA', 'RM 推流接收 + BLE 保持手机/主机连接'],
    ],
    col_widths=[5.5, 2.5, 8]
)

# ═══ 场景3 ═══
heading('四、场景3：BLE 主机 + BLE 从机 + RM 从机', level=2)
para('左耳 GAP_ROLE_ALL（Central 连对侧从机 + Peripheral 广播），同时开启 RM RX。不能 deep sleep。根据 RM、从机、手机的连接状态分为六个子状态。')

make_table(
    ['子状态', '功耗', '备注'],
    [
        ['RM 未连接，无从机，无手机', '1.5 mA', '扫描对侧耳 + 仅 BLE 广播'],
        ['RM 未连接，无从机，有手机', '1.5 mA', '扫描对侧耳 + 手机已连'],
        ['RM 未连接，有从机，有手机', '960 μA', '对侧耳 + 手机双连接，100ms 心跳，停止扫描'],
        ['RM 连接，无从机，无手机', '1.65 mA', 'RM 推流 + 扫描对侧耳 + 仅 BLE 广播'],
        ['RM 连接，无从机，有手机', '1.65 mA', 'RM 推流 + 扫描对侧耳 + 手机已连'],
        ['RM 连接，有从机，有手机', '1.11 mA', 'RM 推流(+150μA) + 对侧耳 + 手机，停止扫描'],
    ],
    col_widths=[5.5, 2.5, 8]
)

# ═══ 场景4 ═══
heading('五、场景4：BLE 主机 + BLE 从机', level=2)
para('左耳 GAP_ROLE_ALL（Central 连对侧从机 + Peripheral 广播），不开启 RM。根据从机和手机的连接状态分为三个子状态。')

make_table(
    ['子状态', '功耗', '备注'],
    [
        ['无从机，无手机', '1.5 mA', '扫描对侧耳 + 仅 BLE 广播'],
        ['无从机，有手机', '1.5 mA', '扫描对侧耳 + 手机已连'],
        ['有从机，有手机', '800 μA', '从机已连，停止扫描，100ms 心跳 + BB_WAKEUP'],
    ],
    col_widths=[5.5, 2.5, 8]
)

# ── Save ──
out = r'd:\projects\onsemi-workspace\docs\charger_box\BODY_POWER_TEST_REPORT.docx'
try:
    doc.save(out)
except PermissionError:
    out = r'd:\projects\onsemi-workspace\docs\charger_box\BODY_POWER_TEST_REPORT_v2.docx'
    doc.save(out)
print(f'Done: {out}')
