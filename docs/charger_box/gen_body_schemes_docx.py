"""Generate body-side architecture schemes Word document — updated power model."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

s = doc.sections[0]
for m in ['top', 'bottom', 'left', 'right']:
    setattr(s, f'{m}_margin', Cm(2))

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text); run.bold = True; run.font.size = Pt(9)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shading = c._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'D9D9D9'
        })
        shading.append(shd)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]; cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text)); run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══ TITLE ═══
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('机身（耳机端）架构方案对比')
r.bold = True; r.font.size = Pt(18)
r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('基于实测功耗数据（2026-08-10）')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(128, 128, 128)
doc.add_paragraph()

# ═══ 前提条件 ═══
heading('一、前提条件与功耗模型', level=2)

heading('1.1 硬件约束', level=3)
make_table(
    ['参数', '值', '说明'],
    [
        ['机身电池', '15mAh', '核心功耗约束'],
        ['BLE 连接距离（机身天线）', '5-11m', '机身天线弱，是系统瓶颈'],
        ['RM TX 距离（充电盒）', '20-56m', '2-4 倍于 BLE'],
        ['充电盒电池', '700mAh', '非瓶颈，续航 42 天'],
    ],
    col_widths=[5, 3, 8]
)

para('核心矛盾：RM 链路覆盖 20-56m，BLE 仅 5-11m。机身 15mAh 电池是唯一硬约束。各场景功耗数据见《机身功耗测试报告》。', bold=True)

# ═══ 方案 A ═══
heading('二、方案 A：纯 Peripheral + BLE 指令型（右耳基线）', level=2)

para('机身仅作为 BLE Peripheral（场景1），充电盒通过 BLE 常连发送 RM_ONOFF 指令切换至 RM RX（场景2）。30s 无 RM 信号自动回退 BLE。对侧耳不直接互联。')

heading('功耗（15mAh 机身）', level=3)
make_table(
    ['模式', '场景', '功耗', '续航'],
    [
        ['BLE 待机', '场景1（手机连接）', '600 μA', '~25h'],
        ['RM 推流', '场景2（RM 连接 + BLE 连接）', '1 mA', '~15h'],
        ['典型日（30% 推流）', '—', '720 μA', '~21h'],
    ],
    col_widths=[4, 5, 3, 3]
)

para('优点：', bold=True)
bullet('功耗最低，待机 600 μA，续航 25h')
bullet('开发量最小，与现有 TX 工程最接近')
bullet('BLE 常连可双向传输音量、电量、校准参数')
bullet('切换延迟 <100ms')

para('缺点：', bold=True)
bullet('覆盖受限于 BLE（5-11m），超出范围后指令失效，依赖 RM 超时兜底')
bullet('两只耳机无直接互联通道')
bullet('充电盒需维持 BLE 多连接，状态机复杂度中等')

para('可实现性：★★★★★ | 充电盒设计文档基准方案。')

# ═══ 方案 B ═══
heading('三、方案 B：双角色主从 + BLE 指令型（当前开发基线）', level=2)

para('左耳 GAP_ROLE_ALL（场景4），右耳 GAP_ROLE_PERIPHERAL（场景1/2）。左耳直连右耳 MAC，100ms 心跳，事件驱动同步（仅程序，不同步音量）。充电盒通过 BLE 指令控制 RM 切换（左耳从场景4→场景3）。')

heading('功耗（15mAh 机身）', level=3)
make_table(
    ['耳侧', '模式', '场景', '功耗', '续航', '说明'],
    [
        ['左耳', 'BLE 待机', '场景4（有从机，有手机）', '800 μA', '~18.75h', '停止扫描 + 100ms 心跳 + BB_WAKEUP'],
        ['左耳', 'RM 推流', '场景3（有从机，有RM）', '1.11 mA', '~13.5h', 'RM初始化+搜索(960) + RM连接(+150)'],
        ['右耳', 'BLE 待机', '场景1（手机连接）', '600 μA', '~25h', '纯 Peripheral + deep sleep'],
        ['右耳', 'RM 推流', '场景2（RM连接+BLE连接）', '1 mA', '~15h', 'RM 推流 + deep sleep'],
    ],
    col_widths=[1.5, 2, 3.5, 2, 2, 5]
)

para('关键发现：从机连上后停止扫描，功耗从 1.5 mA 骤降至 800 μA。RM 初始化（搜索态）增加 160 μA（800→960），RM 连接额外增加 150 μA（960→1.11 mA）。左耳 BLE 待机 ~18.75h，RM 推流 ~13.5h。', bold=True)

heading('双耳同步机制', level=3)
make_table(
    ['同步方向', '触发条件', '机制', '同步内容'],
    [
        ['本机→对侧', '按键 / App 触发', 'GATT Write 对侧 RX [0x01, prog]', '仅程序'],
        ['对侧→本机', '对侧按键触发', 'TX Notify [cnt, prog, vol, 0, 0]', '仅程序'],
        ['连上时', '不触发', '无初始同步', '各自保持当前状态'],
    ],
    col_widths=[3, 3.5, 5.5, 4]
)

para('优点：', bold=True)
bullet('双耳可直接同步程序切换')
bullet('右耳功耗不受影响（600 μA），左耳 960 μA 可接受')
bullet('100ms 心跳双连接稳定，RM + BLE 共存（BBIF_COEX 硬件仲裁）')
bullet('按键推送、广播恢复、ear side 运行时设置等边界 case 已修复')
bullet('BLE 常连保留，充电盒可控、手机 App 可连')

para('缺点：', bold=True)
bullet('左耳 BLE 800 μA / RM 1.11 mA，RM 推流续航 ~13.5h')
bullet('覆盖仍受限于 BLE（5-11m）')
bullet('不能 deep sleep（BB_WAKEUP），是左耳功耗高的根因')
bullet('华为手机 40s 断连未根治（兼容性边界）')
bullet('左右耳需维护两份固件（RM_LEFT vs RM_RIGHT）')

para('可实现性：★★★★☆ | 已实现（DUAL_ROLE.md + SYNC.md），当前开发基线。')

# ═══ 方案 C ═══
heading('四、方案 C：双角色主从 + 按需切换型', level=2)

para('在方案 B 基础上，充电盒不再依赖 BLE 指令——检测到音频直接切 RM TX 广播。机身进入 RM RX 多路径：用户按键 / BLE 指令（范围内）/ App 触发。30s 无 RM 自动回 BLE。')

heading('功耗（同方案 B）', level=3)
make_table(
    ['模式', '场景', '功耗', '续航'],
    [
        ['BLE 待机', '场景4（有从机，有手机）', '800 μA（左耳）', '~18.75h'],
        ['RM 推流', '场景3（有从机，有RM）', '1.11 mA（左耳）', '~13.5h'],
        ['BLE 待机', '场景1（右耳）', '600 μA', '~25h'],
        ['RM 推流', '场景2（右耳）', '1 mA', '~15h'],
    ],
    col_widths=[3, 5, 3.5, 3]
)

para('优点：', bold=True)
bullet('覆盖不受限于 BLE——RM 20-56m 远超 BLE 5-11m')
bullet('充电盒逻辑极简：音频来 → RM TX，无需维护 BLE 连接状态机')
bullet('进入 RM 多路径：按键（不依赖 BLE）+ BLE 指令（范围内最快）')
bullet('保留双耳互联能力')

para('缺点：', bold=True)
bullet('用户需主动按键（或等 BLE 指令），有操作成本')
bullet('按键触发时充电盒可能还没开始 RM TX，机身需短暂等待')
bullet('充电盒不管机身状态，可能"在推流但机身没在听"')
bullet('左耳功耗同方案 B（960 μA）')

para('可实现性：★★★★☆ | 方案 B 的增量演进，风险低。')

# ═══ 方案 D ═══
heading('五、方案 D：双角色主从 + 占空比轮询型（已尝试，已 revert）', level=2)

para('机身 50% 占空比在 BLE 广播（场景1，600 μA）和 RM RX 搜索（场景2，850 μA）之间交替。RM 窗口内检测到信号→锁定推流。充电盒音频来→直接 RM TX。')

heading('功耗（15mAh 机身，右耳视角）', level=3)
make_table(
    ['模式', '功耗', '说明'],
    [
        ['轮询待机（50% 占空比）', '~725 μA', 'BLE 600μA × 50% + RM搜索 850μA × 50%'],
        ['RM 推流', '1 mA', '锁定 RM 后持续推流'],
        ['典型日（30% 推流）', '~808 μA', '70% 轮询 + 30% 推流'],
    ],
    col_widths=[4.5, 3, 8.5]
)

para('注意：此方案仅适用于右耳（纯 Peripheral）。左耳加入后会引入扫描从机的 1.5-1.65 mA，轮询平均功耗急剧上升，不划算。', bold=True)

para('优点：', bold=True)
bullet('覆盖不受限（RM 20-56m）+ 用户无感')
bullet('比 RM 常驻省电，轮询待机 ~725 μA vs 1 mA')
bullet('充电盒逻辑最简')

para('缺点：', bold=True)
bullet('音频启动延迟：最坏 ~2.5s（半个轮询周期）')
bullet('BLE 不保持连接，对侧耳互联需在 BLE 窗口内完成')
bullet('【致命】RM 搜索功耗奇偶交替 bug：RM_Disable 后硬件状态残留，奇数轮正常、偶数轮偏高')
bullet('仅适用于右耳，左耳轮询代价过高')

para('可实现性：★★☆☆☆ | 已实现后 revert（SCHEME4_SLEEP_POLLING.md），需 RM 底层固件配合解决硬件复位问题。')

# ═══ 方案 E ═══
heading('六、方案 E：RM RX 常驻型（续航换覆盖）', level=2)

para('机身始终处于 RM RX 监听模式（场景2，1 mA），不维持 BLE 连接。充电盒音频来→RM TX 广播，机身收到即输出。BLE 仅按需建立（调参/配对）。')

heading('功耗（15mAh 机身）', level=3)
make_table(
    ['模式', '功耗', '续航'],
    [
        ['RM 常驻监听', '1 mA', '~15h'],
        ['RM 推流接收', '1 mA', '~15h'],
    ],
    col_widths=[4, 3, 8]
)

para('优点：', bold=True)
bullet('覆盖最好（RM 20-56m），零切换延迟')
bullet('机身逻辑最简单：有信号就放，没信号就等')
bullet('充电盒逻辑最简')

para('缺点：', bold=True)
bullet('待机续航仅 15h（vs 方案 A 25h，少 40%）')
bullet('无 BLE 常连，音量/参数调节/电量上报需临时建连接')
bullet('对侧耳无法互联')
bullet('15mAh 电池下代价过大')

para('可实现性：★★☆☆☆ | 除非机身电池升级到 20mAh+，否则不推荐。')

# ═══ 七、对比总表 ═══
heading('七、五方案对比总表', level=2)

make_table(
    ['维度', '方案 A\n纯Peri+BLE指令', '方案 B\n双角色+BLE指令', '方案 C\n双角色+按需', '方案 D\n轮询（右耳）', '方案 E\nRM常驻'],
    [
        ['适用耳侧', '右耳', '左耳+右耳', '左耳+右耳', '仅右耳', '通用'],
        ['机身场景', '场景1/2', '左:场景4/3\n右:场景1/2', '同 B', '场景1/2轮询', '场景2常驻'],
        ['待机功耗', '600 μA', '左 800 / 右 600 μA', '同 B', '~725 μA', '1 mA'],
        ['待机续航', '~25h', '左 ~18.75h\n右 ~25h', '同 B', '~20.7h', '~15h'],
        ['RM 推流功耗', '1 mA', '左 1.11 mA\n右 1 mA', '同 B', '1 mA', '1 mA'],
        ['覆盖范围', 'BLE 5-11m', 'BLE 5-11m', 'RM 20-56m', 'RM 20-56m', 'RM 20-56m'],
        ['充电盒复杂度', '中', '中', '最低', '最低', '最低'],
        ['用户操作', '无感', '无感', '需按键', '无感', '无感'],
        ['对侧耳互联', '无', '有（仅程序）', '有（仅程序）', 'BLE窗口内', '无'],
        ['音频启动延迟', '<100ms', '<100ms', '取决于按键', '最坏~2.5s', '零'],
        ['左耳适用', '否', '是', '是', '否（功耗过高）', '是'],
        ['开发状态', '设计完成', '已实现', '未实施', '已revert', '未实施'],
        ['可实现性', '★★★★★', '★★★★☆', '★★★★☆', '★★☆☆☆', '★★☆☆☆'],
    ],
    col_widths=[2.3, 2.5, 2.5, 2.5, 2.5, 2.5]
)

# ═══ 八、推荐路径 ═══
heading('八、推荐路径', level=2)

heading('短期：方案 B（当前基线）', level=3)
bullet('双耳互联 + BLE 指令控制已实现，覆盖核心需求')
bullet('左耳 960 μA 可接受，续航 ~15.6h')
bullet('右耳 600 μA 表现优秀，续航 ~25h')

heading('中期：方案 B → 方案 C 增量演进', level=3)
bullet('在方案 B 基础上加按键触发 RM RX 入口 + 充电盒简化状态机')
bullet('覆盖从 5-11m 提升到 20-56m')
bullet('增量改动小，风险低')

heading('长期：方案 D（需先解决 RM 硬件复位问题）', level=3)
bullet('用户体验最好（无感 + 覆盖好 + 充电盒最简）')
bullet('但仅适用于右耳，左耳扫描功耗 1.5 mA 让轮询无意义')
bullet('需 RM 底层固件配合解决功耗奇偶交替 bug')

heading('方案 E：不推荐', level=3)
para('40% 续航损失（25h→15h）在 15mAh 硬约束下代价过大。')

# ═══ 九、决策树 ═══
heading('九、方案选择决策树', level=2)
para('按以下优先级判断：', bold=True)
doc.add_paragraph()
bullet('Q1: 是否需要双耳直接互联？')
para('    否 → 方案 A（最简单，600 μA，25h）')
para('    是 → 继续 Q2')
bullet('Q2: 用户能否接受按键操作？')
para('    能 → 方案 C（覆盖好 + 保留双耳互联，充电盒最简）')
para('    不能（必须无感）→ 继续 Q3')
bullet('Q3: 能否接受 BLE 距离限制（5-11m）？')
para('    能 → 方案 B（当前基线，已实现）')
para('    不能 → 方案 D（需先解决 RM 硬件复位问题，且仅适用于右耳）')

# ═══ 参考文档 ═══
heading('十、参考文档', level=2)
make_table(
    ['文档', '内容'],
    [
        ['docs/charger_box/BODY_POWER_TEST_REPORT.docx', '机身功耗测试报告（三种场景 × 子状态）'],
        ['docs/charger_box/BODY_CHARGER_SCHEMES.txt', '机身&充电盒配合四方案对比（原始数据）'],
        ['docs/charger_box/CHARGER_BOX_DESIGN.txt', '充电盒方案设计书（硬件架构、状态机）'],
        ['docs/peer_ear/DUAL_ROLE.md', '双角色主从连接（GAP_ROLE_ALL、100ms心跳、功耗实测）'],
        ['docs/peer_ear/SYNC.md', '双耳同步方案（程序同步、防回环）'],
        ['docs/sleep/SCHEME4_SLEEP_POLLING.md', '方案四轮询实现 + RM功耗bug分析'],
        ['docs/rm_connection_plan.md', 'RM 连接功能集成方案'],
        ['docs/开发/硬件配置.md', '硬件配置（RF功率、时钟、供电、低功耗条件）'],
    ],
    col_widths=[5.5, 10.5]
)

# ── Save ──
out = r'd:\projects\onsemi-workspace\docs\charger_box\BODY_ARCHITECTURE_SCHEMES.docx'
try:
    doc.save(out)
except PermissionError:
    out = r'd:\projects\onsemi-workspace\docs\charger_box\BODY_ARCHITECTURE_SCHEMES_v2.docx'
    doc.save(out)
print(f'Done: {out}')
