"""Generate body-charger scheme comparison Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
s = doc.sections[0]
for m in ['top','bottom','left','right']:
    setattr(s, f'{m}_margin', Cm(2))
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text); run.bold = True; run.font.size = Pt(9)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shading = c._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'):'clear', qn('w:color'):'auto', qn('w:fill'):'D9D9D9'
        })
        shading.append(shd)
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = t.rows[r+1].cells[c]; cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text)); run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══ TITLE ═══
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('机身 & 充电盒 配合方案对比')
r.bold = True; r.font.size = Pt(18)
r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('2026-08-05')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(128,128,128)
doc.add_paragraph()

# ═══ Prerequisites ═══
heading('前提条件（来自实测数据）', level=2)
make_table(
    ['参数', '值', '说明'],
    [
        ['BLE 连接距离（机身）', '5-11m', '机身天线弱，是瓶颈'],
        ['RM TX 距离（充电盒）', '20-56m', '2-4 倍于 BLE'],
        ['机身 BLE 从机功耗', '600μA', '持续广播 + 连接'],
        ['机身 RM RX 功耗', '900μA', '持续监听（无信号时）'],
        ['机身 RM RX 推流功耗', '~1mA', '接收并输出音频时'],
    ],
    col_widths=[4.5, 3.5, 8]
)

heading('充电盒功耗（实测）', level=3)
make_table(
    ['芯片', '模式', '功耗'],
    [
        ['RSL10', 'BLE 主机连接机身', '200μA'],
        ['AD6976', '连接手机（待机）', '~500μA'],
        ['AD6976', '播放音频', '6mA'],
        ['RFX2401C PA', 'TX 使能', '60mA'],
    ],
    col_widths=[4, 7, 5]
)

make_table(
    ['状态', 'RSL10', 'AD6976', 'PA', '合计', '续航(700mAh)'],
    [
        ['BLE 待机', '200μA', '500μA', '0', '~0.7mA', '~1000h（42天）'],
        ['音频推流（RM TX）', '200μA', '6mA', '60mA', '~66mA', '~10.5h'],
    ],
    col_widths=[3.5, 1.8, 2, 1.8, 2.2, 3.5]
)

para('充电盒核心功能：充电（座舱大电池供电）+ 控制（RSL10 主控）。', bold=True)
bullet('机身放入充电盒 → 充电 + RSL10 待机，AD6976 关闭')
bullet('机身拿出 → RSL10 检测 → 开启 AD6976 → BLE 扫描/连接')
bullet('充电盒 BLE 待机仅 0.7mA，续航 42 天，功耗不是瓶颈')

para('核心矛盾：RM 链路覆盖 20-56m，BLE 仅 5-11m。机身 15mAh 电池是唯一约束。', bold=True)

# ═══ Scheme 1 ═══
heading('方案一：BLE 指令型', level=2)
para('机身默认 BLE 从机（600μA），充电盒 BLE 常连。音频来 → BLE 发 RM_ONOFF=1 → 机身切 RM RX；音频停 → BLE 发 RM_ONOFF=0 → 机身切回 BLE。机身 RM RX 模式 30s 无 RM 信号 → 自动回 BLE（安全兜底）。')

heading('功耗估算（15mAh 机身 / 700mAh 充电盒）', level=3)
make_table(
    ['使用模式', '机身功耗', '充电盒功耗', '机身续航', '充电盒续航'],
    [
        ['BLE 待机', '600μA', '0.7mA', '~25h', '~1000h'],
        ['RM 推流', '900μA', '66mA', '~17h', '~10.5h'],
        ['典型日（30% 推流）', '690μA', '~20mA', '~22h', '~35h'],
    ],
    col_widths=[3.8, 2.5, 2.5, 2.5, 2.5]
)

para('优点：', bold=True)
bullet('BLE 指令控制，切换快（<100ms），充电盒明确知道机身状态')
bullet('BLE 常连可传音量、电量、校准参数')
bullet('可渐进优化：长 BLE interval 降待机至 ~500μA（续航 30h）')

para('缺点：', bold=True)
bullet('覆盖受限于 BLE 距离（5-11m），超出范围指令失效，依赖机身 RM 超时兜底')

para('可实现性：★★★★★ — 与现有 TX 工程最接近，改动量最小。')

# ═══ Scheme 2 ═══
heading('方案二：RM RX 常驻型', level=2)
para('机身始终 RM RX（900μA），不维持 BLE 连接。前提：BLE 模式下无法检测 RM 载波，必须 RM RX 才能监听。充电盒检测到音频 → 直接切 RM TX 广播。机身检测到 RM → 音频输出；信号消失 → 继续 RM RX 等待。BLE 仅按需建立（配对、调音量）。')

heading('功耗估算（15mAh 机身 / 700mAh 充电盒）', level=3)
make_table(
    ['使用模式', '机身功耗', '充电盒功耗', '机身续航', '充电盒续航'],
    [
        ['RM RX 常驻', '900μA', '0.7mA', '~17h', '~1000h'],
        ['RM 推流接收', '~1mA', '66mA', '~15h', '~10.5h'],
    ],
    col_widths=[3.8, 2.5, 2.5, 2.5, 2.5]
)

para('优点：', bold=True)
bullet('覆盖不受限：RM 20-56m，远超 BLE 5-11m')
bullet('机身逻辑最简单：有信号就放，没信号就等')
bullet('零切换延迟，天然不怕 BLE 断连')

para('缺点：', bold=True)
bullet('机身持续 900μA-1mA，15mAh 仅 15-17h（比方案一少 ~30%）')
bullet('无 BLE 常连，音量/参数调节需临时建连接')

para('可实现性：★★☆☆☆ — 机身需 RM RX 常驻 + 信号检测逻辑。')

# ═══ Scheme 3 ═══
heading('方案三：按需切换型', level=2)
para('机身默认 BLE 从机（600μA）低功耗待机。进入 RM：用户按键 或 充电盒发 BLE 指令 → 机身切 RM RX。退出 RM：RM RX 模式 30s 无 RM → 自动切回 BLE。充电盒：检测到音频 → 直接切 RM TX 广播（不等机身确认，逻辑极简）。')

heading('功耗估算（15mAh 机身 / 700mAh 充电盒）', level=3)
make_table(
    ['使用模式', '机身功耗', '充电盒功耗', '机身续航', '充电盒续航'],
    [
        ['BLE 待机（默认）', '600μA', '0.7mA', '~25h', '~1000h'],
        ['RM 推流接收', '~1mA', '66mA', '~15h', '~10.5h'],
        ['典型日（30% 推流）', '720μA', '~20mA', '~21h', '~35h'],
    ],
    col_widths=[3.8, 2.5, 2.5, 2.5, 2.5]
)

para('优点：', bold=True)
bullet('默认 BLE 低功耗，续航与方案一持平（~25h）')
bullet('进入 RM 多路径：按键（不依赖 BLE）+ BLE 指令（快速）')
bullet('充电盒不管机身状态，逻辑极简：音频来 → RM TX')
bullet('超出 BLE 范围时，用户可按机身按键主动切 RM')
bullet('30s 无 RM 自动回 BLE，不怕忘记切回')

para('缺点：', bold=True)
bullet('用户需主动按键（或等 BLE 指令），有操作成本')
bullet('按键触发时充电盒可能还没开始 RM TX，机身需短暂等待')

para('可实现性：★★★★☆ — 在方案一基础上加按键触发 + 充电盒去掉 BLE 指令依赖。')

# ═══ Scheme 4 ═══
heading('方案四：占空比轮询型', level=2)
para('机身不维持 BLE 连接，仅 BLE 广播。50% 占空比在 BLE（600μA）和 RM RX（900μA）之间交替，平均功耗 750μA。在 RM RX 窗口内检测到 RM 信号 → 锁定 RM RX 接收音频。RM 信号消失 30s → 恢复轮询。充电盒：检测到音频 → 直接切 RM TX 广播。')

heading('功耗估算（15mAh 机身 / 700mAh 充电盒）', level=3)
make_table(
    ['使用模式', '机身功耗', '充电盒功耗', '机身续航', '充电盒续航'],
    [
        ['轮询待机（50% 占空比）', '750μA', '0.7mA', '~20h', '~1000h'],
        ['RM 推流接收', '~1mA', '66mA', '~15h', '~10.5h'],
    ],
    col_widths=[3.8, 2.5, 2.5, 2.5, 2.5]
)

para('优点：', bold=True)
bullet('不依赖 BLE 连接，覆盖范围同 RM（20-56m）')
bullet('比方案二省电 ~17%（900→750μA），续航多 3h')
bullet('无需用户操作：自动轮询，RM 窗口内检测到自动切换')
bullet('充电盒逻辑极简：音频来 → RM TX')

para('缺点：', bold=True)
bullet('音频启动延迟：最坏情况需等一个轮询周期（BLE 窗口内收不到 RM）')
bullet('BLE 不保持连接，音量/参数调节需等 BLE 窗口或手动触发')
bullet('轮询逻辑增加机身状态机复杂度')

para('可实现性：★★★☆☆ — 机身需轮询调度 + RM 信号检测。充电盒改动小。')

# ═══ Comparison ═══
heading('四方案对比总表', level=2)
make_table(
    ['维度', '方案一', '方案二', '方案三', '方案四'],
    [
        ['机身默认模式', 'BLE 从机 600μA', 'RM RX 900μA', 'BLE 从机 600μA', 'BLE/RM 50%轮询'],
        ['进入 RM 方式', 'BLE 指令', '始终在 RM RX', '按键 / BLE 指令', 'RM 窗口自动检测'],
        ['退出 RM 方式', 'BLE 指令/30s超时', '始终在 RM RX', '30s 无 RM 自动回', '30s 无 RM 回轮询'],
        ['充电盒逻辑', '音频→BLE指令→RM TX', '音频→直接RM TX', '音频→直接RM TX', '音频→直接RM TX'],
        ['覆盖范围', '受限于 BLE(5-11m)', '不受限(20-56m)', '不受限(20-56m)', '不受限(20-56m)'],
        ['机身待机功耗', '600μA', '900μA', '600μA', '750μA'],
        ['机身续航（待机）', '~25h', '~17h', '~25h', '~20h'],
        ['机身续航（典型日）', '~22h', '~17h', '~21h', '~19h'],
        ['音频启动延迟', '<100ms', '零延迟', '取决于按键时机', '最坏一个轮询周期'],
        ['机身逻辑复杂度', '低', '最低', '中', '较高'],
        ['充电盒逻辑复杂度', '中', '低', '最低', '最低'],
        ['用户操作', '无感', '无感', '需按键/等指令', '无感'],
        ['安全兜底', 'RM 30s 超时', '天然机制', 'RM 30s 超时', 'RM 30s 超时'],
    ],
    col_widths=[2.8, 3, 3, 3, 3]
)

# ═══ Recommendation ═══
heading('推荐路径', level=2)

make_table(
    ['维度', '方案一', '方案二', '方案三', '方案四'],
    [
        ['续航', '25h ★★', '17h ★', '25h ★★', '20h ★★'],
        ['覆盖范围', '5-11m ★', '20-56m ★★★', '20-56m ★★★', '20-56m ★★★'],
        ['用户体验', '无感 ★★★', '无感 ★★★', '需按键 ★★', '无感 ★★★'],
        ['充电盒复杂度', '中', '低', '最低', '最低'],
        ['开发难度', '低 ★★★', '高 ★', '中 ★★', '中 ★★'],
        ['音频启动延迟', '快', '零', '取决于按键', '最坏~1周期'],
    ],
    col_widths=[2.8, 3, 3, 3, 3]
)

para('推荐：方案一先行（续航好、开发快），方案四作为目标（覆盖好、充电盒最简、用户无感）。方案三适用于需用户自主控制的场景。方案二续航牺牲 30%，如非必要不选。', bold=True)

out = r'd:\projects\onsemi-workspace\docs\BODY_CHARGER_SCHEMES.docx'
try:
    doc.save(out)
except PermissionError:
    out = r'd:\projects\onsemi-workspace\docs\BODY_CHARGER_SCHEMES_v2.docx'
    doc.save(out)
print(f'Done: {out}')
