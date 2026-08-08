"""Generate scan test report Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shading = hdr_cells[i]._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'D9D9D9'
        })
        shading.append(shd)
    for r, row in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, text in enumerate(row):
            row_cells[c].text = ''
            p = row_cells[c].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ═══ TITLE ═══
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RFX2401C 扫描/TX 距离测试报告')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('测试日期：2026-08-04    扫描程序：TX 主机扫描程序（纯扫描，不连接）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ═══ 一、测试环境 ═══
heading('一、测试环境', level=2)

heading('扫描端（烧录 TX 程序）', level=3)
make_table(
    ['编号', '硬件', 'RFX2401C', '备注'],
    [
        ['S1', 'RSL10 开发板', '无', '内部 radio'],
        ['S2', 'Smart1604 充电盒蓝牙空板', '有', 'IO10=RXEN, IO11=TXEN，扫描时开 RX 使能'],
    ],
    col_widths=[1.5, 5.5, 2.5, 6.5]
)

heading('广播端（sleep 设备）', level=3)
make_table(
    ['编号', '硬件'],
    [
        ['B1', 'Smart1604 机身'],
        ['B2', 'Smart1654 机身'],
        ['B3', 'Smart1654 开发板'],
    ],
    col_widths=[2, 14]
)

# ═══ 二、扫描程序配置 ═══
heading('二、扫描程序配置', level=2)
bullet('IO10 → RXEN（LNA enable），IO11 → TXEN（PA enable）')
bullet('开机默认 RX 模式（RXEN=HIGH, TXEN=LOW），持续 BLE 主动扫描')
bullet('匹配方式：MAC 地址匹配（6 个 peer）')
bullet('扫描超时后自动重启扫描')

# ═══ 三、TX 功率参数 ═══
heading('三、TX 功率参数', level=2)
make_table(
    ['位置', '参数', '值'],
    [
        ['app.h:65', 'OUTPUT_POWER_6DBM', '1（启用）'],
        ['app_init.c:402', 'Sys_RFFE_SetTXPower(6)', 'BLE 初始化后设 6dBm'],
        ['app_process.c:208', 'Sys_RFFE_SetTXPower(6)', '进入 RM 推流时设 6dBm'],
        ['app_init.c:73', 'VDDPA_ENABLE', 'VDDPA_DISABLE_BITBAND（内部 PA 关）'],
        ['app_init.c:74', 'VDDPA_SW_CTRL', 'VDDPA_SW_VDDRF_BITBAND'],
    ],
    col_widths=[4, 5.5, 6.5]
)
para('注：S2（带 PA）RSL10 内部 TX 功率 6dBm 作为 RFX2401C PA 输入。S1（无 PA）内部 radio 直接输出 6dBm。')

# ═══ 四、扫描测试结果 ═══
heading('四、扫描测试结果（BLE 主机扫描）', level=2)
make_table(
    ['扫描端', '广播端', '最大可扫描距离'],
    [
        ['S1 RSL10 开发板（无 PA）', 'B1 Smart1604 机身', '~8m'],
        ['S1 RSL10 开发板（无 PA）', 'B2 Smart1654 机身', '~5m'],
        ['S1 RSL10 开发板（无 PA）', 'B3 Smart1654 开发板', '~20m'],
        ['S2 Smart1604 空板（有 PA，RX 使能）', 'B1 Smart1604 机身', '~11m'],
        ['S2 Smart1604 空板（有 PA，RX 使能）', 'B2 Smart1654 机身', '~8m'],
        ['S2 Smart1604 空板（有 PA，RX 使能）', 'B3 Smart1654 开发板', '~23m'],
    ],
    col_widths=[7.5, 4.5, 4]
)
bullet('纯扫描模式，不发起连接')
bullet('扫描到设备后仅打印 MAC，不停止扫描')

# ═══ 五、TX 发送测试结果 ═══
heading('五、TX 发送测试结果（RM TX 推流）', level=2)
para('发送端配置：烧录 TX 程序，开机默认 RM TX 模式（TXEN=HIGH, RXEN=LOW）。接收端：sleep 设备。', bold=False)

make_table(
    ['发送端', '接收端', '条件', '结果'],
    [
        ['S1 RSL10 开发板（无 PA）', 'B3 Smart1654 开发板', '穿墙，预研办公室', '~15m'],
        ['S1 RSL10 开发板（无 PA）', 'B3 Smart1654 开发板', '空旷场地', '~56m'],
        ['S1 RSL10 开发板（无 PA）', 'B1 Smart1604 机身', '办公区无遮挡', '~10m'],
        ['S1 RSL10 开发板（无 PA）', 'B1 Smart1604 机身', '空旷场地', '~22m'],
        ['S1 RSL10 开发板（无 PA）', 'B2 Smart1654 机身', '办公区无遮挡', '~10m'],
        ['S1 RSL10 开发板（无 PA）', 'B2 Smart1654 机身', '空旷场地', '~20m'],
        ['S2 Smart1604 空板（有 PA，TX 使能）', 'B1 Smart1604 机身', '穿墙，研发办公室→厕所', '通过'],
        ['S2 Smart1604 空板（有 PA，TX 使能）', 'B2 Smart1654 机身', '穿墙，研发办公室→厕所', '通过'],
        ['S2 Smart1604 空板（有 PA，TX 使能）', 'B3 Smart1654 开发板', '穿墙，研发办公室→厕所', '通过'],
    ],
    col_widths=[6.5, 4, 4, 1.5]
)
para('注：S1 无 PA 发送时 Smart1604/Smart1654 机身不支持穿墙；S2 有 PA 时三个设备均支持穿墙。')

out = r'd:\projects\onsemi-workspace\docs\RFX2401C_SCAN_TEST.docx'
doc.save(out)
print(f'Done: {out}')
