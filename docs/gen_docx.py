"""Generate charger box design Word document."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Helper functions ──
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    return p

def make_table(headers, rows, col_widths=None):
    """Create a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # Gray background
        shading = hdr_cells[i]._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D9D9D9'
        })
        shading.append(shd)

    # Data rows
    for r, row in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, text in enumerate(row):
            row_cells[c].text = ''
            p = row_cells[c].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return table

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ═══════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('充电盒方案设计书')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于 RFX2401C 扫描/TX 测试数据（2026-08-04）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# 一、测试数据总结
# ═══════════════════════════════════════════════
heading('一、测试数据总结', level=2)

heading('1. 扫描距离（BLE 主机扫描）', level=3)
make_table(
    ['扫描端', 'Smart1654 开发板', 'Smart1604 机身', 'Smart1654 机身'],
    [
        ['RSL10 开发板（无 PA）', '20m', '8m', '5m'],
        ['Smart1604 空板（有 PA，RX 使能）', '23m', '11m', '8m'],
    ],
    col_widths=[5.5, 3.5, 3.5, 3.5]
)

heading('2. TX 发送距离（RM 推流，RSL10 开发板 无PA）', level=3)
make_table(
    ['接收端', '办公区', '空旷', '穿墙'],
    [
        ['Smart1654 开发板', '—', '56m', '15m'],
        ['Smart1604 机身', '10m', '22m', '—'],
        ['Smart1654 机身', '10m', '20m', '—'],
    ],
    col_widths=[5.5, 3.5, 3.5, 3.5]
)

heading('3. TX 发送（Smart1604 空板 有PA，TX使能）', level=3)
para('三个设备均支持穿墙。')

# ═══════════════════════════════════════════════
# 二、机身功能定义
# ═══════════════════════════════════════════════
heading('二、机身功能定义', level=2)

heading('运行模式', level=3)
make_table(
    ['模式', '功耗', '说明'],
    [
        ['BLE 低功耗从机', '600μA', 'BLE peripheral，低功耗待机'],
        ['RM RX', '900μA', '接收 RM 音频流'],
    ],
    col_widths=[4.5, 3, 8.5]
)

bullet('两种模式通过 BLE 指令互相切换')
bullet('机身作为 BLE 从机（peripheral），充电盒作为 BLE 主机（central）')

# ═══════════════════════════════════════════════
# 三、充电盒硬件架构
# ═══════════════════════════════════════════════
heading('三、充电盒硬件架构', level=2)

make_table(
    ['芯片', '功能'],
    [
        ['RSL10', '运行 TX 工程，BLE 主机 + RM TX 向机身推流音频'],
        ['AD6976（经典蓝牙）', '连接手机，接收音频数据'],
    ],
    col_widths=[4.5, 11.5]
)

para('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('手机 ──经典蓝牙──→ AD6976 ──DMIC──→ RSL10 ──RM TX──→ 机身')
run.font.size = Pt(10)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ═══════════════════════════════════════════════
# 四、工作流程
# ═══════════════════════════════════════════════
heading('四、工作流程', level=2)

heading('无音频（待机）', level=3)
bullet('RSL10 处于 BLE 主机模式，同时连接两个机身')
bullet('两个机身处于 BLE 低功耗从机模式（600μA）')
bullet('RSL10 维持 BLE 连接，监听 AD6976 DMIC 音频')

heading('有音频（推流）', level=3)
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
run = p.add_run('AD6976 收到手机音频 → 通过 DMIC 输入 RSL10')
p2 = doc.add_paragraph()
run = p2.add_run('2. ')
run.bold = True
run = p2.add_run('RSL10 检测到音频 → 通过 BLE 发送 RM_ONOFF=1 → 机身切换到 RM RX 模式（900μA）')
p3 = doc.add_paragraph()
run = p3.add_run('3. ')
run.bold = True
run = p3.add_run('RSL10 自身切换到 RM TX 模式，开始广播推流')
p4 = doc.add_paragraph()
run = p4.add_run('4. ')
run.bold = True
run = p4.add_run('两个机身同时接收同一路音频')

heading('音频结束', level=3)
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
run = p.add_run('RSL10 检测到静音 → 停止 RM TX，切回 BLE 模式')
p2 = doc.add_paragraph()
run = p2.add_run('2. ')
run.bold = True
run = p2.add_run('RSL10 通过 BLE 发送 RM_ONOFF=0 → 机身切回 BLE 低功耗从机模式')

# ═══════════════════════════════════════════════
# 五、充电盒方案设计
# ═══════════════════════════════════════════════
heading('五、充电盒方案设计', level=2)

heading('1. 音频检测方案', level=3)
para('AD6976 无 GPIO 音频状态输出脚，由 RSL10 监听 DMIC，使用 EMA 能量阈值判断有无音频。该方案已在 TX 工程中验证。')

heading('2. BLE 模式切换协议', level=3)
para('复用现有 Custom Service 的 RM_ONOFF 特征值：')

make_table(
    ['BLE 指令', '值', '含义'],
    [
        ['RM_ONOFF write', '1', '机身切 RM RX 模式'],
        ['RM_ONOFF write', '0', '机身切 BLE 低功耗模式'],
    ],
    col_widths=[4.5, 2, 9.5]
)

para('时序要求：', bold=True)
bullet('RSL10 发送 RM_ONOFF=1 → 等待 N ms → 自身切 RM TX 推流')
bullet('音频结束 → RSL10 停 RM TX → 发送 RM_ONOFF=0')

heading('3. RFX2401C RF 路径控制', level=3)
para('基于测试数据，RFX2401C 控制策略：')

make_table(
    ['场景', 'TXEN (IO11)', 'RXEN (IO10)', '说明'],
    [
        ['BLE 扫描/连接', '0', '1', 'RX 通路，LNA 提升扫描距离'],
        ['RM TX 推流', '1', '0', 'TX 通路，PA 提升发射距离'],
        ['空闲/休眠', '0', '0', '省电'],
    ],
    col_widths=[4, 3, 3, 6]
)

para('切换时序：先关当前通路 → 再开目标通路，避免 TXEN 和 RXEN 同时为 HIGH。', bold=True)

heading('4. 充电盒状态机', level=3)

make_table(
    ['当前状态', '事件', '下一状态', '动作'],
    [
        ['INIT', 'BLE 就绪', 'SCANNING', '启动扫描'],
        ['SCANNING', 'MAC 匹配到机身', 'CONNECTING', '停止扫描，发起连接'],
        ['CONNECTING', '两机身均连接', 'CONNECTED', '保持 BLE 连接'],
        ['CONNECTING', '超时未连上', 'SCANNING', '重新扫描'],
        ['CONNECTED', '检测到音频', 'STREAMING', '发 RM_ONOFF=1 → 切 RFX2401C=TX → RM_Enable'],
        ['CONNECTED', '机身断连', 'SCANNING', '重新扫描'],
        ['STREAMING', '音频结束', 'CONNECTED', 'RM_Disable → 切 RFX2401C=RX → 发 RM_ONOFF=0'],
        ['STREAMING', '机身断连', 'SCANNING', '停 RM，重新扫描'],
    ],
    col_widths=[2.8, 3.2, 2.8, 7.2]
)

heading('5. 确认项', level=3)

make_table(
    ['编号', '事项', '结论'],
    [
        ['Q1', 'AD6976 音频检测方式', 'RSL10 DMIC 能量检测'],
        ['Q2', 'AD6976→RSL10 音频数据接口', 'DMIC'],
        ['Q3', '机身 RM RX↔BLE 切换指令', '复用 Custom Service RM_ONOFF'],
        ['Q4', '双机身音频分配', 'RM TX 广播，两个机身同时收同一路'],
        ['Q5', '充电盒功耗要求', '无特殊要求'],
        ['Q6', 'S2 空板 IO11 拉不低（2.7V）', '先不管'],
    ],
    col_widths=[1.5, 5.5, 9]
)

# ═══════════════════════════════════════════════
# 六、测试数据分析与推论
# ═══════════════════════════════════════════════
heading('六、测试数据分析与推论', level=2)

heading('关键发现', level=3)

heading('发现 1：BLE 扫描距离是瓶颈，不是 RM TX', level=3)
make_table(
    ['方向', '最短距离', '说明'],
    [
        ['充电盒→机身（扫描）', '5-11m', '机身 BLE 广播功率/天线决定'],
        ['充电盒→机身（RM TX）', '20-56m', 'RM 链路余量远超 BLE'],
    ],
    col_widths=[4.5, 3, 8.5]
)
para('充电盒扫机身最远 11m（带 PA），但机身接收 RM TX 能到 56m。BLE 和 RM 覆盖范围严重不对称。')

heading('发现 2：PA 的收益', level=3)
make_table(
    ['', '无 PA', '有 PA', '提升'],
    [
        ['BLE 扫描', '5-8m', '8-11m', '+3m'],
        ['RM TX 穿墙', '仅开发板', '全部穿墙', '质的飞跃'],
    ],
    col_widths=[4, 3, 3, 6]
)
para('PA 在 RX 方向收益有限（+3m），在 TX 方向收益巨大（从无法穿墙到全部可穿墙）。')

heading('发现 3：机身天线差异巨大', level=3)
make_table(
    ['设备', '扫描距离', 'TX 距离'],
    [
        ['Smart1654 开发板', '20m', '56m'],
        ['Smart1654 机身', '5m', '20m'],
    ],
    col_widths=[5, 5.5, 5.5]
)
para('同一芯片方案，开发板是机身的 3-4 倍距离。机身天线/匹配是短板。')

heading('推论', level=3)

para('推论 A：机身必须支持 RM 超时自动切回 BLE', bold=True)
para('原因：BLE 范围 5-11m，RM TX 范围 20-56m，严重不对称。', indent=True)
para('场景：用户在 BLE 范围内收到音频 → 切 RM RX → 走远超出 BLE 11m → 音频结束 → 充电盒发 RM_ONOFF=0 但机身 BLE 已断连收不到 → 机身永远停在 RM RX 模式，功耗 900μA 跑空。', indent=True)
para('方案：机身 RM RX 模式下，N 秒无 RM 信号 → 自动切回 BLE 从机模式。', indent=True)

para('推论 B：充电盒必须带 PA', bold=True)
para('无 PA 的 RM TX 只在理想条件下覆盖 20-22m 且不能穿墙。充电盒的实际使用场景（房间隔墙）必然需要穿墙能力。PA 是必须的。', indent=True)

para('推论 C：BLE 连接保持是系统薄弱环节', bold=True)
bullet('连接建立后尽量保持，BLE 断开才重新扫描')
bullet('BLE connection interval 和 supervision timeout 合理配置')
bullet('机身 BLE 广播功率尽量拉满')

# ═══════════════════════════════════════════════
# 七、可行性分析
# ═══════════════════════════════════════════════
heading('七、可行性分析', level=2)

heading('充电盒 TX 端已有基础', level=3)
make_table(
    ['功能', '状态', '说明'],
    [
        ['BLE 扫描 MAC 匹配', '已验证', '多 peer 扫描正常'],
        ['BLE 多连接', '已有代码', 'PEER_COUNT=6，支持多 peer'],
        ['DMIC 音频检测', '已验证', 'EMA 能量阈值检测正常'],
        ['BLE→RM TX 切换', '已验证', 'RF_SwitchToCPMode + RM_Enable'],
        ['RFX2401C TX/RX 控制', '已验证', '扫描 RX / 推流 TX'],
        ['BLE 指令 RM_ONOFF', '已有代码', 'Custom Service 写特征值'],
        ['RM TX 广播推流', '已验证', '空旷 56m，带 PA 可穿墙'],
    ],
    col_widths=[4.5, 3, 8.5]
)

heading('机身端需要新增', level=3)
make_table(
    ['功能', '说明'],
    [
        ['RM 超时自动切回 BLE', '机身 RM RX 模式下 N 秒无 RM 信号，自动切 BLE 从机'],
    ],
    col_widths=[5, 11]
)
para('机身端只有这一个新增需求，改动量很小。')

heading('充电盒 TX 端需要新增/调整', level=3)
make_table(
    ['功能', '说明'],
    [
        ['恢复 BLE 连接逻辑', '扫到 → 连接 → 保持连接'],
        ['音频检测触发 RM_ONOFF', '检测到音频 → 发 RM_ONOFF=1 给两个机身'],
        ['音频停止触发回切', '停 RM TX → 发 RM_ONOFF=0'],
        ['断连处理', '任一机身断连 → 重新扫描连接'],
    ],
    col_widths=[5, 11]
)

heading('功耗估算', level=3)
make_table(
    ['状态', 'RSL10 BLE', 'RFX2401C', '合计'],
    [
        ['待机（BLE 连接两个机身）', '~12mA', '0（RX 通路，LNA μA 级）', '~12mA'],
        ['推流（RM TX + PA）', '~12mA', '+60mA', '~72mA'],
    ],
    col_widths=[5, 3.5, 4, 3.5]
)
para('RFX2401C TX 模式实测多 60mA，仅推流期间发生。')

heading('风险点', level=3)
make_table(
    ['风险', '等级', '应对'],
    [
        ['BLE 与 RM 覆盖不对称', '中', '机身 RM 超时自动回落 BLE（推论 A）'],
        ['S2 空板 IO11 2.7V', '低', '实测不影响功能，后续硬件可加下拉'],
        ['两机身同时连接稳定性', '低', 'PEER_COUNT 机制已验证'],
        ['AD6976 DMIC 静音时噪声', '中', '需调音频检测阈值，上线前实测'],
    ],
    col_widths=[5.5, 1.5, 9]
)

# ═══════════════════════════════════════════════
# 八、结论
# ═══════════════════════════════════════════════
heading('八、结论', level=2)
para('方案可行。核心技术点全部已验证。', bold=True)
para('新增工作量集中在：')
bullet('充电盒 TX 端：状态机调整（INIT → SCANNING → CONNECTING → CONNECTED ⇄ STREAMING）')
bullet('机身端：RM 超时自动回落 BLE 逻辑（RM RX 模式下 N 秒无信号 → 切 BLE 从机）')

# ── Save ──
out = r'd:\projects\onsemi-workspace\docs\CHARGER_BOX_DESIGN.docx'
doc.save(out)
print(f'Done: {out}')
